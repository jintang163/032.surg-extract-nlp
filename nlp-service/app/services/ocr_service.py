import os
import re
import io
from loguru import logger
from typing import Optional, Tuple
from pathlib import Path


class OcrDependencyError(Exception):
    def __init__(self, message: str, detail: str = ""):
        super().__init__(message)
        self.detail = detail


class OcrProcessingError(Exception):
    def __init__(self, message: str, stage: str = "unknown", detail: str = ""):
        super().__init__(message)
        self.stage = stage
        self.detail = detail


try:
    from PIL import Image
except ImportError:
    Image = None

try:
    import pytesseract
except ImportError:
    pytesseract = None

try:
    from pdf2image import convert_from_bytes, convert_from_path
except ImportError:
    convert_from_bytes = None
    convert_from_path = None

try:
    import pdfplumber
except ImportError:
    pdfplumber = None

try:
    from docx import Document
except ImportError:
    Document = None

from app.config import get_settings


class OcrService:
    def __init__(self):
        self.settings = get_settings()
        self._ensure_dirs()
        self._setup_tesseract()
        self._check_dependencies()

    def _ensure_dirs(self):
        for d in [self.settings.upload_dir, self.settings.temp_dir, self.settings.output_dir]:
            Path(d).mkdir(parents=True, exist_ok=True)

    def _setup_tesseract(self):
        if pytesseract is not None:
            try:
                pytesseract.pytesseract.tesseract_cmd = self.settings.tesseract_cmd
            except Exception as e:
                logger.warning(f"Tesseract配置失败: {e}")

    def _check_dependencies(self):
        missing = []
        if pytesseract is None:
            missing.append("pytesseract")
        if Image is None:
            missing.append("Pillow")
        if missing:
            logger.warning(
                f"以下OCR相关依赖未安装，OCR功能不可用: {', '.join(missing)}. "
                f"请执行: pip install pytesseract Pillow pdf2image pdfplumber python-docx "
                f"并安装Tesseract-OCR引擎"
            )

    def _ensure_ocr_dependencies(self):
        missing = []
        if pytesseract is None:
            missing.append("pytesseract")
        if Image is None:
            missing.append("Pillow")
        if missing:
            raise OcrDependencyError(
                f"OCR功能依赖未安装: {', '.join(missing)}",
                detail="请执行 pip install pytesseract Pillow 并安装 Tesseract-OCR 引擎"
            )

    def process_file(
        self, file_content: bytes, filename: str, file_type: Optional[str] = None
    ) -> Tuple[str, str]:
        logger.info(f"开始处理文件: {filename}, type={file_type}")

        if file_content is None or len(file_content) == 0:
            raise OcrProcessingError(
                "文件内容为空",
                stage="input_validation",
                detail="上传的文件没有内容，请检查文件是否损坏"
            )

        if not file_type:
            file_type = self._detect_file_type(filename)

        try:
            raw_text = self._extract_text(file_content, filename, file_type)
        except OcrProcessingError:
            raise
        except OcrDependencyError:
            raise
        except Exception as e:
            logger.error(f"文本提取失败: {filename}, {e}", exc_info=True)
            raise OcrProcessingError(
                f"文本提取失败: {str(e)}",
                stage="text_extraction",
                detail=str(e)
            )

        if not raw_text or len(raw_text.strip()) == 0:
            raise OcrProcessingError(
                "未识别到有效文本内容",
                stage="text_extraction",
                detail="OCR识别结果为空，请上传更清晰的文档扫描件或使用可编辑的Word/PDF文件"
            )

        try:
            processed_text = self._preprocess_text(raw_text)
        except Exception as e:
            logger.warning(f"文本预处理失败，使用原始文本: {e}")
            processed_text = raw_text

        logger.info(
            f"OCR处理完成: 原始长度={len(raw_text)}, 处理后长度={len(processed_text)}"
        )
        return raw_text, processed_text

    def process_file_by_path(
        self, file_path: str, file_type: Optional[str] = None
    ) -> Tuple[str, str]:
        logger.info(f"通过文件路径处理: {file_path}, type={file_type}")

        if not file_path or not os.path.exists(file_path):
            raise OcrProcessingError(
                f"文件不存在: {file_path}",
                stage="file_validation",
                detail=f"无法在路径 {file_path} 找到文件"
            )

        try:
            file_size = os.path.getsize(file_path)
            if file_size == 0:
                raise OcrProcessingError(
                    "文件大小为0字节",
                    stage="file_validation",
                    detail="上传的文件为空，请检查"
                )
            if file_size > 100 * 1024 * 1024:
                raise OcrProcessingError(
                    f"文件过大: {file_size / 1024 / 1024:.2f}MB，最大支持100MB",
                    stage="file_validation"
                )
        except OcrProcessingError:
            raise
        except Exception as e:
            raise OcrProcessingError(
                f"读取文件元信息失败: {str(e)}",
                stage="file_validation",
                detail=str(e)
            )

        filename = os.path.basename(file_path)
        if not file_type:
            file_type = self._detect_file_type(filename)

        try:
            with open(file_path, "rb") as f:
                file_content = f.read()
        except Exception as e:
            raise OcrProcessingError(
                f"读取文件失败: {str(e)}",
                stage="file_read",
                detail=str(e)
            )

        return self.process_file(file_content, filename, file_type)

    def _extract_text(self, file_content: bytes, filename: str, file_type: str) -> str:
        file_type = (file_type or "").upper()

        if file_type == "TEXT" or filename.lower().endswith(".txt"):
            return self._extract_from_text(file_content)

        if file_type == "WORD" or filename.lower().endswith((".doc", ".docx")):
            return self._extract_from_docx(file_content, filename)

        if file_type == "PDF" or filename.lower().endswith(".pdf"):
            return self._extract_from_pdf(file_content, filename)

        if file_type == "IMAGE" or filename.lower().endswith(
            (".png", ".jpg", ".jpeg", ".gif", ".bmp", ".tiff", ".tif")
        ):
            return self._extract_from_image(file_content, filename)

        raise OcrProcessingError(
            f"不支持的文件类型: {file_type} (文件: {filename})",
            stage="file_type_detection",
            detail="支持的格式: txt, doc, docx, pdf, png, jpg, jpeg, gif, bmp, tiff"
        )

    def _detect_file_type(self, filename: str) -> str:
        ext = filename.split(".")[-1].lower() if "." in filename else ""
        if ext == "txt":
            return "TEXT"
        if ext in ("doc", "docx"):
            return "WORD"
        if ext == "pdf":
            return "PDF"
        if ext in ("png", "jpg", "jpeg", "gif", "bmp", "tiff", "tif"):
            return "IMAGE"
        return "UNKNOWN"

    def _extract_from_text(self, file_content: bytes) -> str:
        encoding_candidates = ["utf-8", "utf-8-sig", "gbk", "gb2312", "gb18030", "latin-1"]
        last_error = None
        for encoding in encoding_candidates:
            try:
                text = file_content.decode(encoding)
                logger.debug(f"使用编码 {encoding} 解析TEXT文件成功")
                return text
            except UnicodeDecodeError as e:
                last_error = e
                continue
        raise OcrProcessingError(
            "无法识别文本文件编码",
            stage="text_decode",
            detail=f"尝试的编码: {encoding_candidates}, 最后错误: {last_error}"
        )

    def _extract_from_docx(self, file_content: bytes, filename: str) -> str:
        if Document is None:
            raise OcrDependencyError(
                "python-docx 未安装，无法解析Word文档",
                detail="请执行: pip install python-docx"
            )

        try:
            doc = Document(io.BytesIO(file_content))
        except Exception as e:
            raise OcrProcessingError(
                f"Word文档解析失败: {str(e)}",
                stage="docx_parse",
                detail="可能是文件格式不兼容或文件已损坏，请尝试转换为其他格式"
            )

        paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
        tables_text = []
        try:
            for table in doc.tables:
                for row in table.rows:
                    row_text = [cell.text.strip() for cell in row.cells]
                    tables_text.append(" | ".join(row_text))
        except Exception as e:
            logger.warning(f"提取Word表格失败，跳过表格内容: {e}")

        result = "\n".join(paragraphs + tables_text)
        if not result.strip():
            logger.warning(f"Word文档 {filename} 未提取到任何文本内容")
        return result

    def _extract_from_pdf(self, file_content: bytes, filename: str) -> str:
        texts = []

        if pdfplumber is not None:
            try:
                with pdfplumber.open(io.BytesIO(file_content)) as pdf:
                    for idx, page in enumerate(pdf.pages):
                        try:
                            page_text = page.extract_text()
                            if page_text and page_text.strip():
                                texts.append(page_text.strip())
                            else:
                                logger.debug(f"PDF第{idx + 1}页未提取到文本，将尝试OCR")
                        except Exception as e:
                            logger.warning(f"PDF第{idx + 1}页文本提取失败: {e}")
            except Exception as e:
                logger.warning(f"PDF文本层提取失败，将使用OCR: {e}")

        if texts and sum(len(t.strip()) for t in texts) >= 50:
            logger.info(f"PDF文本层提取成功: {len(texts)}页")
            return "\n\n".join(texts)

        self._ensure_ocr_dependencies()
        if convert_from_bytes is None:
            raise OcrDependencyError(
                "pdf2image 未安装，无法将PDF转为图片进行OCR",
                detail="请执行: pip install pdf2image 并安装 poppler"
            )

        try:
            images = convert_from_bytes(file_content, dpi=300)
            logger.info(f"PDF转图片OCR: 共{len(images)}页")
        except Exception as e:
            raise OcrProcessingError(
                f"PDF转图片失败: {str(e)}",
                stage="pdf_to_image",
                detail="请确认已安装Poppler，或上传更小的PDF文件"
            )

        ocr_results = []
        for idx, img in enumerate(images):
            try:
                page_text = self._ocr_image(img)
                if page_text and page_text.strip():
                    ocr_results.append(page_text.strip())
            except Exception as e:
                logger.warning(f"PDF第{idx + 1}页OCR失败: {e}")
                continue

        if not ocr_results:
            raise OcrProcessingError(
                "PDF OCR识别失败，未提取到有效文本",
                stage="pdf_ocr",
                detail="请确认PDF文件为清晰的扫描件，或上传可编辑的PDF版本"
            )

        return "\n\n".join(ocr_results)

    def _extract_from_image(self, file_content: bytes, filename: str) -> str:
        self._ensure_ocr_dependencies()

        try:
            img = Image.open(io.BytesIO(file_content))
            img.verify()
            img = Image.open(io.BytesIO(file_content))
        except Exception as e:
            raise OcrProcessingError(
                f"图片加载失败: {str(e)}",
                stage="image_load",
                detail="图片文件可能已损坏，或格式不受支持"
            )

        if img.mode != "RGB":
            try:
                img = img.convert("RGB")
            except Exception as e:
                logger.warning(f"图片格式转换失败，继续尝试OCR: {e}")

        try:
            text = self._ocr_image(img)
        except Exception as e:
            raise OcrProcessingError(
                f"图片OCR识别失败: {str(e)}",
                stage="image_ocr",
                detail=str(e)
            )

        if not text or len(text.strip()) < 5:
            raise OcrProcessingError(
                "图片OCR未识别到足够的文字内容",
                stage="image_ocr",
                detail="请上传更清晰的手术记录扫描件或照片，确保文字可辨认"
            )

        return text

    def _ocr_image(self, image) -> str:
        self._ensure_ocr_dependencies()

        try:
            config = "--psm 6"
            text = pytesseract.image_to_string(
                image,
                lang=self.settings.tesseract_lang,
                config=config,
            )
        except pytesseract.TesseractNotFoundError:
            raise OcrDependencyError(
                "Tesseract-OCR 引擎未安装或路径未配置",
                detail="请安装Tesseract-OCR，并设置环境变量或在配置中指定 tesseract_cmd 路径"
            )
        except pytesseract.TesseractError as e:
            raise OcrProcessingError(
                f"Tesseract OCR 执行失败: {e.message}",
                stage="tesseract_exec",
                detail=f"status={e.status}, message={e.message}"
            )
        except Exception as e:
            raise OcrProcessingError(
                f"OCR识别异常: {str(e)}",
                stage="tesseract_exec",
                detail=str(e)
            )

        if not isinstance(text, str):
            text = str(text)

        return text

    def _preprocess_text(self, text: str) -> str:
        if not text:
            return ""

        lines = text.split("\n")
        cleaned_lines = []
        for line in lines:
            line = line.strip()
            if not line:
                cleaned_lines.append("")
                continue
            line = self._remove_noise(line)
            if line:
                cleaned_lines.append(line)

        cleaned_text = "\n".join(cleaned_lines)
        cleaned_text = re.sub(r"\n{3,}", "\n\n", cleaned_text)
        cleaned_text = re.sub(r"[ \t]{2,}", " ", cleaned_text)
        cleaned_text = re.sub(r"\u3000", " ", cleaned_text)

        return cleaned_text.strip()

    def _remove_noise(self, line: str) -> str:
        noise_patterns = [
            (r"^[-_=\*·•]{3,}$", ""),
            (r"第\s*\d+\s*页\s*[共/｜丨/|/\\]\s*\d+\s*页", ""),
            (r"第\s*\d+\s*/\s*\d+\s*页", ""),
            (r"Page\s*\d+\s*of\s*\d+", "", True),
            (r"P\.?\s*\d+\s*[/／]\s*\d+", "", True),
            (r"打印时间\s*[:：].*", ""),
            (r"录入时间\s*[:：].*", ""),
            (r"审核时间\s*[:：].*", ""),
            (r"记录时间\s*[:：].*", ""),
            (r"医院名称\s*[:：].*", ""),
            (r"病历号\s*[:：].*", "", True),
            (r"^[\|│┃┄┅━┃┆┇┊┋]{2,}", ""),
            (r"[\|│┃┄┅━┃┆┇┊┋]", " "),
            (r"\s{2,}", " "),
        ]

        for item in noise_patterns:
            pattern = item[0]
            repl = item[1]
            ignore_case = item[2] if len(item) > 2 else False
            flags = re.IGNORECASE if ignore_case else 0
            try:
                line = re.sub(pattern, repl, line, flags=flags)
            except re.error:
                continue

        return line.strip()
